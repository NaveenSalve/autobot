from __future__ import annotations
import io, json, re
from typing import Any

from api.services.ai_router import generate_text


def extract_text_from_upload(filename: str, content: bytes) -> str:
    """Extract text from PDF/DOCX/TXT upload with graceful fallback."""
    lower = filename.lower()
    if lower.endswith('.pdf'):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype='pdf')
            return '\n'.join(page.get_text('text') for page in doc)[:80_000]
        except Exception as exc:
            raise ValueError(f'Could not parse PDF: {exc}') from exc
    if lower.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(' | '.join(cell.text for cell in row.cells))
            return '\n'.join(paragraphs + tables)[:80_000]
        except Exception as exc:
            raise ValueError(f'Could not parse DOCX: {exc}') from exc
    if lower.endswith(('.txt', '.md', '.rtf')):
        return content.decode('utf-8', errors='ignore')[:80_000]
    raise ValueError('Unsupported file type. Upload PDF, DOCX, TXT, MD, or RTF.')


def _json_from_text(text: str) -> dict[str, Any] | None:
    try:
        return json.loads(text)
    except Exception:
        pass
    match = re.search(r'\{.*\}', text, flags=re.S)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            return None
    return None


def heuristic_profile(text: str) -> dict[str, Any]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    joined = '\n'.join(lines)
    email = next(iter(re.findall(r'[\w.%-]+@[\w.-]+\.[A-Za-z]{2,}', joined)), '')
    name = lines[0][:80] if lines else ''
    headline = ''
    for line in lines[1:8]:
        if len(line) > 12 and not re.search(r'@|\+?\d[\d\s-]{7,}', line):
            headline = line[:160]
            break
    service_keywords = [
        'web design', 'website development', 'react', 'frontend', 'backend', 'seo', 'automation',
        'ai chatbot', 'chatbot', 'lead generation', 'crm', 'email marketing', 'copywriting', 'analytics'
    ]
    low = joined.lower()
    services = []
    for key in service_keywords:
        if key in low:
            label = 'AI chatbot setup' if key == 'chatbot' else key.title()
            if label not in services:
                services.append(label)
    if not services:
        services = ['Website audit', 'AI automation', 'Lead generation']
    industries = []
    for key in ['hospitality', 'healthcare', 'real estate', 'fitness', 'restaurant', 'retail', 'saas', 'education']:
        if key in low:
            industries.append(key.title())
    return {
        'name': name,
        'email': email,
        'headline': headline or 'Business development consultant',
        'bio': joined[:1200],
        'services': services[:8],
        'target_industries': industries[:8],
        'target_locations': [],
    }


async def extract_profile_with_ai(text: str) -> dict[str, Any]:
    fallback = heuristic_profile(text)
    prompt = f"""
Extract a sales profile from this resume/profile text. Return ONLY valid JSON with keys:
name, email, headline, bio, services, target_industries, target_locations.
Arrays must contain strings. Keep bio under 900 characters.

TEXT:
{text[:12000]}
"""
    try:
        raw = await generate_text(prompt, system='You extract structured JSON from resumes. Output only JSON.', temperature=0.1, max_tokens=1200)
        parsed = _json_from_text(raw)
        if not parsed:
            return fallback
        return {
            'name': str(parsed.get('name') or fallback['name']),
            'email': str(parsed.get('email') or fallback['email']),
            'headline': str(parsed.get('headline') or fallback['headline']),
            'bio': str(parsed.get('bio') or fallback['bio'])[:1200],
            'services': parsed.get('services') if isinstance(parsed.get('services'), list) else fallback['services'],
            'target_industries': parsed.get('target_industries') if isinstance(parsed.get('target_industries'), list) else fallback['target_industries'],
            'target_locations': parsed.get('target_locations') if isinstance(parsed.get('target_locations'), list) else fallback['target_locations'],
        }
    except Exception:
        return fallback
